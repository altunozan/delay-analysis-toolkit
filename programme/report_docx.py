"""Module 11 — Report Assembler.

Assembles the per-module narratives and key findings into one Word document:
title page, one chapter per included module, a single deduplicated
Limitations section aggregated from every module's caveats and warnings, and
a "Basis of Analysis" appendix recording exactly what the report was built
from (files, SHA-256 hashes, data dates, settings used).

The aggregation is pure concatenation/deduplication — no LLM. Narrative
chapters reproduce the module narratives the analyst generated (and
reviewed) in the app; modules without a generated narrative contribute their
deterministic key findings only.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x38, 0x64)

STANDING_REPORT_CAVEAT = (
    "This document is a preliminary factual screening assembled from the "
    "programme files listed in the Basis of Analysis appendix. It describes "
    "movement and change recorded in those files; it does not attribute "
    "cause or responsibility and is not a cause-linked delay analysis."
)


@dataclass
class ReportSection:
    title: str
    narrative_md: str | None = None       # analyst-generated narrative
    key_findings: list[str] = field(default_factory=list)
    caveats: list[str] = field(default_factory=list)   # feeds Limitations


@dataclass
class SourceFile:
    file_name: str
    sha256: str
    data_date: datetime | None
    role: str                             # Baseline / Update / Current
    activity_count: int


@dataclass
class BasisOfAnalysis:
    files: list[SourceFile] = field(default_factory=list)
    settings: list[str] = field(default_factory=list)  # "Module — setting: x"
    generated_at: datetime = field(default_factory=datetime.now)
    tool_note: str = ("Deterministic engines computed all figures; where an "
                      "AI narrative is included it was generated from those "
                      "figures under fixed objectivity rules and reviewed "
                      "by the analyst before inclusion.")


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_md_paragraph(doc: Document, text: str, style: str | None = None):
    """Add one paragraph, converting **bold** spans."""
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def _add_markdown(doc: Document, md: str, base_heading_level: int) -> None:
    """Minimal markdown renderer: #/##/### headings, - bullets, paragraphs."""
    for raw in md.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            hashes, text = m.groups()
            # Narrative "##" becomes a level under the chapter heading.
            level = min(base_heading_level + max(len(hashes) - 2, 0), 4)
            doc.add_heading(_BOLD_RE.sub(r"\1", text), level=level)
        elif re.match(r"^\s*[-*•]\s+", line):
            _add_md_paragraph(doc, re.sub(r"^\s*[-*•]\s+", "", line),
                              style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            _add_md_paragraph(doc, re.sub(r"^\s*\d+\.\s+", "", line),
                              style="List Number")
        else:
            _add_md_paragraph(doc, line)


def _dedupe(items: list[str]) -> list[str]:
    seen, out = set(), []
    for i in items:
        key = i.strip().lower()
        if key and key not in seen:
            seen.add(key)
            out.append(i.strip())
    return out


def build_assembled_report(
    report_title: str,
    project_name: str,
    author: str,
    sections: list[ReportSection],
    basis: BasisOfAnalysis,
) -> bytes:
    """Assemble the full preliminary report as a .docx (returns bytes)."""
    doc = Document()
    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(10.5)

    # --- title page -------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(report_title or "Preliminary Delay Analysis Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    for line in (project_name, author,
                 f"Issued {basis.generated_at:%d %B %Y}",
                 "PRELIMINARY — for review and discussion"):
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line).font.size = Pt(12)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(STANDING_REPORT_CAVEAT).italic = True
    doc.add_page_break()

    # --- contents (static list; Word can replace with a live TOC) ---------
    doc.add_heading("Contents", level=1)
    for i, s in enumerate(sections, start=1):
        doc.add_paragraph(f"{i}. {s.title}", style="List Number")
    doc.add_paragraph(f"{len(sections) + 1}. Limitations",
                      style="List Number")
    doc.add_paragraph(f"{len(sections) + 2}. Appendix A — Basis of Analysis",
                      style="List Number")
    doc.add_page_break()

    # --- module chapters ---------------------------------------------------
    for i, s in enumerate(sections, start=1):
        doc.add_heading(f"{i}. {s.title}", level=1)
        if s.key_findings:
            doc.add_heading("Key figures", level=2)
            for kf in s.key_findings:
                _add_md_paragraph(doc, kf, style="List Bullet")
        if s.narrative_md:
            _add_markdown(doc, s.narrative_md, base_heading_level=2)
        else:
            p = doc.add_paragraph()
            p.add_run(
                "No narrative was generated for this module; the key "
                "figures above are reported without commentary."
            ).italic = True

    # --- limitations (aggregated, deduplicated) ----------------------------
    doc.add_heading(f"{len(sections) + 1}. Limitations", level=1)
    doc.add_paragraph(
        "The following limitations apply to this report. They are "
        "aggregated from every analysis included and are reproduced in "
        "full; none has been omitted."
    )
    all_caveats = _dedupe(
        [STANDING_REPORT_CAVEAT]
        + [c for s in sections for c in s.caveats])
    for c in all_caveats:
        _add_md_paragraph(doc, c, style="List Bullet")

    # --- basis of analysis --------------------------------------------------
    doc.add_page_break()
    doc.add_heading(f"{len(sections) + 2}. Appendix A — Basis of Analysis",
                    level=1)
    doc.add_paragraph(
        "This appendix records the source files, their integrity hashes, "
        "and the analysis settings from which every figure in this report "
        "was computed, so the analysis can be independently reproduced."
    )
    doc.add_heading("Source programme files", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(["File", "Role", "Data date", "Activities",
                           "SHA-256 (first 16)"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for f in basis.files:
        row = table.add_row().cells
        row[0].text = f.file_name
        row[1].text = f.role
        row[2].text = (f"{f.data_date:%Y-%m-%d}" if f.data_date else "—")
        row[3].text = str(f.activity_count)
        row[4].text = f.sha256[:16]

    if basis.settings:
        doc.add_heading("Analysis settings", level=2)
        for s in basis.settings:
            _add_md_paragraph(doc, s, style="List Bullet")

    doc.add_heading("Method note", level=2)
    doc.add_paragraph(basis.tool_note)
    doc.add_paragraph(
        f"Report assembled {basis.generated_at:%Y-%m-%d %H:%M}."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
