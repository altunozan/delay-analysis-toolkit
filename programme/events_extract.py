"""Event extraction from correspondence and event narratives (TIA intake).

Reads letters / instructions / notices (txt, docx, pdf) or a pasted dated
narrative, and builds the prompt for an AI extraction pass that proposes
CANDIDATE delay events. The forensic rails are structural:

- Every candidate must cite its source document AND a VERBATIM snippet.
- The snippet is verified against the actual document text here, in code —
  a candidate whose quoted evidence cannot be found in its claimed source
  is DROPPED and counted, never shown as evidence.
- Candidates are proposals for the analyst's event form; nothing enters
  the register without the analyst adopting it.

Aligned with AACE RP 52R-06 (Prospective Time Impact Analysis): events
extracted here feed a TIA performed against the last accepted schedule
update prior to the event, with the fragnet kept to the fewest activities
practical.

No API calls in this module — prompt builders and strict parsers only.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from datetime import datetime

MAX_DOC_CHARS = 16_000          # per document, keeps prompts bounded
MAX_TOTAL_CHARS = 60_000

EXTRACTION_SYSTEM_PROMPT = (
    "You are a construction claims analyst extracting potential delay "
    "events from project correspondence. You return ONLY valid JSON — no "
    "commentary, no fences. Every event MUST cite its source document by "
    "name and quote a VERBATIM snippet (copied exactly, character for "
    "character) from that document as evidence. You never invent events, "
    "dates, or quotations: if the documents describe no delay event, "
    "return an empty list."
)


# --------------------------------------------------------------------------- #
# Document reading (txt / docx / pdf) — heavy imports stay lazy
# --------------------------------------------------------------------------- #

def read_document(file_name: str, raw: bytes) -> str:
    """Extract plain text from an uploaded letter. Empty string on failure."""
    lower = file_name.lower()
    try:
        if lower.endswith(".docx"):
            import io as _io

            from docx import Document
            doc = Document(_io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(p for p in parts if p and p.strip())
        if lower.endswith(".pdf"):
            import io as _io

            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(raw))
            return "\n".join((page.extract_text() or "")
                             for page in reader.pages)
        return raw.decode("utf-8", errors="replace")
    except Exception:                                    # noqa: BLE001
        return ""


# --------------------------------------------------------------------------- #
# Candidate model + prompt
# --------------------------------------------------------------------------- #

@dataclass
class EventCandidate:
    title: str
    description: str = ""
    date_start: datetime | None = None
    other_dates: list[str] = field(default_factory=list)
    party_asserted: str = ""
    affected_scope: str = ""
    source_doc: str = ""
    source_snippet: str = ""
    confidence: str = "medium"
    verified: bool = False          # snippet found verbatim in the source


def build_event_extraction_prompt(
    docs: list[tuple[str, str]],
) -> str:
    """docs: (name, text). Pasted narratives arrive as a doc too."""
    lines = [
        "<task>Extract every POTENTIAL delay event described in these "
        "project documents: employer instructions, variations, late "
        "drawings/approvals, procurement delays, access restrictions, "
        "utility clashes, weather, suspensions, rework, acceleration, "
        "testing failures, or similar. One entry per distinct event. "
        "These are candidates for a prospective Time Impact Analysis "
        "(AACE RP 52R-06), so capture the dates and instructed/delayed "
        "scope precisely.</task>",
        "",
        "<documents>",
    ]
    total = 0
    for name, text in docs:
        clipped = (text or "")[:MAX_DOC_CHARS]
        if total + len(clipped) > MAX_TOTAL_CHARS:
            clipped = clipped[:max(MAX_TOTAL_CHARS - total, 0)]
        total += len(clipped)
        lines.append(f"<doc name={json.dumps(name)}>")
        lines.append(clipped)
        lines.append("</doc>")
    lines += [
        "</documents>",
        "",
        '<output>Return ONLY JSON: {"events": [{"title": "...", '
        '"description": "what happened / what was instructed", '
        '"date_start": "YYYY-MM-DD or null", '
        '"other_dates": ["YYYY-MM-DD", ...], '
        '"party_asserted": "who the document asserts is responsible '
        '(or empty)", "affected_scope": "area/discipline/work package", '
        '"source_doc": "exact document name", '
        '"source_snippet": "VERBATIM quotation from that document, max '
        '240 characters, copied exactly", '
        '"confidence": "low|medium|high"}]}. '
        "An event without a verbatim snippet must not be returned. If no "
        'delay events are described, return {"events": []}.</output>',
    ]
    return "\n".join(lines)


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip().lower()


def parse_event_candidates(
    text: str, docs: list[tuple[str, str]],
) -> tuple[list[EventCandidate], int]:
    """Parse + VERIFY. Returns (verified candidates, dropped count).

    A candidate is dropped when its claimed source document was not
    supplied, or its quoted snippet does not occur verbatim (whitespace-
    normalised) in that document — the structural defence against
    fabricated evidence.
    """
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```[a-zA-Z]*\n?", "", cleaned)
        cleaned = re.sub(r"\n?```\s*$", "", cleaned)
    start, end = cleaned.find("{"), cleaned.rfind("}")
    if start == -1 or end <= start:
        return [], 0
    try:
        obj = json.loads(cleaned[start:end + 1])
    except json.JSONDecodeError:
        return [], 0
    items = obj.get("events") if isinstance(obj, dict) else None
    if not isinstance(items, list):
        return [], 0

    doc_text = {name: _norm(t) for name, t in docs}
    out: list[EventCandidate] = []
    dropped = 0
    for i in items:
        if not isinstance(i, dict):
            continue
        src = str(i.get("source_doc", "")).strip()
        snippet = str(i.get("source_snippet", "")).strip()[:300]
        title = str(i.get("title", "")).strip()
        if not title or not src or not snippet:
            dropped += 1
            continue
        if src not in doc_text or _norm(snippet) not in doc_text[src]:
            dropped += 1                 # unverifiable evidence -> gone
            continue
        try:
            ds = i.get("date_start")
            date_start = (datetime.strptime(str(ds), "%Y-%m-%d")
                          if ds else None)
        except ValueError:
            date_start = None
        conf = str(i.get("confidence", "medium")).lower()
        out.append(EventCandidate(
            title=title[:160],
            description=str(i.get("description", ""))[:600],
            date_start=date_start,
            other_dates=[str(d) for d in (i.get("other_dates") or [])
                         if re.match(r"^\d{4}-\d{2}-\d{2}$", str(d))][:8],
            party_asserted=str(i.get("party_asserted", ""))[:120],
            affected_scope=str(i.get("affected_scope", ""))[:200],
            source_doc=src, source_snippet=snippet,
            confidence=conf if conf in ("low", "medium", "high")
            else "medium",
            verified=True))
    return out, dropped


# --------------------------------------------------------------------------- #
# AACE RP 52R-06 — analysis-schedule selection
# --------------------------------------------------------------------------- #

RP52R06_CAVEAT = (
    "Performed as a prospective Time Impact Analysis in line with AACE RP "
    "52R-06: the fragnet models the event with the fewest activities "
    "practical, is inserted into the most recent accepted schedule update "
    "prior to the event, and the time impact is the difference between "
    "the pre- and post-insertion completion forecasts. Excusability and "
    "compensability are contractual determinations outside this "
    "calculation."
)


def recommended_analysis_schedule(
    revisions_meta: list[tuple[str, datetime | None]],
    event_date: datetime | None,
) -> str | None:
    """52R-06: the last update whose data date precedes the event.

    ``revisions_meta`` — (label, data_date) in data-date order. Returns
    the recommended label, or None when it cannot be determined.
    """
    if event_date is None:
        return None
    best = None
    for label, dd in revisions_meta:
        if dd is not None and dd <= event_date:
            best = label
    return best
